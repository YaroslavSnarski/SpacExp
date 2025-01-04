import platform
import subprocess
from .base_processor import FileProcessor
import os
import time
from docx import Document  # обработка docx без Word

# условный импорт для Windows
if platform.system() == "Windows":
    try:
        import pythoncom
        from win32com.client import Dispatch
        WIN32_AVAILABLE = True
    except ImportError:
        WIN32_AVAILABLE = False
else:
    WIN32_AVAILABLE = False

from .logging_config import setup_logging
process_logger, error_logger = setup_logging(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

class BaseDOCXProcessor(FileProcessor):
    """
    Класс для обработки файлов формата DOC и DOCX, извлекающий метаданные: автор, заголовок и количество страниц.

    Наследует FileProcessor и предоставляет методы для работы с документами Microsoft Word.
    """
    def __init__(self, default_author="Unknown Author"):
        """
        Инициализирует процессор DOCX с указанным автором по умолчанию.

        Аргументы:
            default_author (str): Имя автора, используемое по умолчанию, если не указано иное.
        """
        super().__init__(default_author=default_author)

    def is_word_installed(self):
        """
        Проверяет, установлен ли Microsoft Word на текущей системе.

        Возвращает:
            bool: True, если Word установлен, иначе False.
        """
        if platform.system() == "Windows" and WIN32_AVAILABLE:
            try:
                pythoncom.CoInitialize()
                word = Dispatch("Word.Application")
                word.Quit()
                return True
            except Exception:
                return False
            finally:
                pythoncom.CoUninitialize()
        return False

    def count_word_pages(self, filepath):
        """
        Подсчитывает количество страниц в файле DOC или DOCX в зависимости от системы и наличия Word.

        Аргументы:
            file_path (str): Путь к файлу DOC или DOCX.

        Возвращает:
            int: Количество страниц в документе или None, если подсчет невозможен.
        """
        filepath = os.path.abspath(filepath)
        if not os.path.exists(filepath):
            error_logger.error(f"File not found: {filepath}")
            return None

        system = platform.system()

        if system == "Windows" and WIN32_AVAILABLE:
            if self.is_word_installed():
                return self._count_pages_with_word(filepath)
            else:
                return self._count_pages_with_docx(filepath)
        elif system == "Linux":
            if filepath.endswith('.doc'):
                docx_filepath = self.convert_doc_to_docx(filepath)
                print(docx_filepath)
                return self._count_pages_with_docx(docx_filepath)
            else:
                return self._count_pages_with_docx(filepath)
        else:
            error_logger.error(f"Unsupported system for page count: {system}")
            return None

    def _count_pages_with_word(self, filepath):
        """
        Подсчет страниц через Microsoft Word.

        Аргументы:
            filepath (str): Путь к файлу.

        Возвращает:
            int: Количество страниц или None при ошибке.
        """
        if not WIN32_AVAILABLE:
            error_logger.error("win32com is not available on this system.")
            return None

        pythoncom.CoInitialize()
        word = None
        try:
            word = Dispatch("Word.Application")
            doc = word.Documents.Open(filepath, ReadOnly=True)
            page_count = doc.ComputeStatistics(2)  # wdStatisticPages
            process_logger.info(f"Calculated page count for {filepath}: {page_count}")
            doc.Close(False)
            return page_count
        except Exception as e:
            error_logger.error(f"Failed to estimate pages for {filepath} with Word: {e}")
            return None
        finally:
            if word:
                word.Quit()
            pythoncom.CoUninitialize()

    def _count_pages_with_docx(self, filepath):
        """
        Подсчет страниц с использованием библиотеки python-docx.

        Аргументы:
            filepath (str): Путь к файлу.

        Возвращает:
            int: Оценочное количество страниц или None при ошибке.
        """
        try:
            doc = Document(filepath)
            #paragraphs = len(doc.paragraphs)
            # количество символов в документе
            total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            # примерная оценка, которая может сильно отклоняться
            #estimated_pages = max(1, paragraphs // 7)  # пусть 7 параграфов на страницу
            estimated_pages = max(1, total_chars // 1300)  # пусть 1.3k символов на страницу
            process_logger.info(f"Estimated page count for {filepath} using python-docx: {estimated_pages}")
            return estimated_pages
        except Exception as e:
            error_logger.error(f"Failed to estimate pages for {filepath} with python-docx: {e}")
            return None

    def get_generic_doc_info(self, filepath, default_author=None):
        """
        Основной метод для извлечения общей информации о документе.

        Аргументы:
            filepath (str): Путь к файлу.

        Возвращает:
            dict: Словарь с общей информацией о файле и метаданными документа.
        """
        page_count = self.count_word_pages(filepath)
        return {
            "author": default_author or self.default_author,
            "page_count": page_count,
        }

    def process_document(self, filepath):
        """Основной метод для обработки DOCX и DOC файлов."""
        file_info = self.get_generic_info(filepath)
        if not file_info:
            return None

        extension = file_info.get("extension")
        if extension in {"docx", "doc"}:
            doc_info = self.get_generic_doc_info(filepath)
            file_info.update(doc_info)
        else:
            error_logger.error(f"Unsupported file type: {filepath}")
            return None

        process_logger.info(f"Successfully processed: {filepath}")
        return file_info


class DOCXProcessor(BaseDOCXProcessor):
    """Класс для обработки DOCX и DOC файлов локально."""
    def process(self, filepath):
        """Обрабатывает файл локально."""
        return self.process_document(filepath)


class DOCXProcessorWeb(BaseDOCXProcessor):
    """
    Класс для обработки DOCX и DOC файлов в веб-приложении, извлекает информацию о количестве страниц и типе файла.

    Наследует BaseDOCXProcessor и предоставляет методы для работы с документами Microsoft Word в контексте веб-приложений.
    """
    def process(self, filepath):
        """
        Обрабатывает файл DOCX или DOC и возвращает информацию, включая количество страниц.

        Аргументы:
            file_path (str): Путь к файлу, который необходимо обработать.

        Возвращает:
            dict: Словарь с информацией о файле, включая количество страниц и тип.
        """
        try:
            start_time = time.time()
            file_info = self.process_document(filepath)
            elapsed_time = time.time() - start_time
            process_logger.info(f"DOCX processed: {filepath} in {elapsed_time:.2f} seconds")
            return file_info
        except Exception as e:
            error_logger.error(f"Error processing DOCX {filepath}: {e}")
            return {"type": "document", "error": str(e)}


    def convert_doc_to_docx(self, doc_filepath):
        """
        Конвертирует DOC файл в DOCX с помощью LibreOffice на Linux.
        """
        try:
            # Определим путь для выходного файла
            docx_filepath = os.path.splitext(doc_filepath)[0] + '.docx'
            
            # Выполним команду для конвертации через LibreOffice
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', os.path.dirname(doc_filepath), doc_filepath], check=True)            
            process_logger.info(f"Converted {doc_filepath} to {docx_filepath}")
            return docx_filepath
        except subprocess.CalledProcessError as e:
            error_logger.error(f"Failed to convert {doc_filepath} to DOCX: {e}")
            return None
